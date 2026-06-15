"""
多格式文档加载器
支持 PDF / Word / Excel / CSV / Markdown / TXT 等常见格式
"""

from pathlib import Path

import pandas as pd

from app.core.logging import log


class DocumentLoaderError(Exception):
    """文档加载异常"""

    pass


class DocumentLoader:
    """
    多格式文档加载器
    将各种文件格式统一提取为纯文本
    """

    # 支持的最大文件大小：20MB
    MAX_FILE_SIZE = 20 * 1024 * 1024

    def __init__(self, max_file_size: int = MAX_FILE_SIZE):
        self.max_file_size = max_file_size

    def load(self, file_path: str | Path) -> str:
        """
        加载文件并返回文本内容

        Args:
            file_path: 文件路径

        Returns:
            提取的文本内容
        """
        path = Path(file_path)

        if not path.exists():
            raise DocumentLoaderError(f"文件不存在: {path}")

        if not path.is_file():
            raise DocumentLoaderError(f"路径不是文件: {path}")

        if path.stat().st_size > self.max_file_size:
            raise DocumentLoaderError(
                f"文件过大: {path.name} ({path.stat().st_size / 1024 / 1024:.2f} MB), "
                f"最大支持 {self.max_file_size / 1024 / 1024:.0f} MB"
            )

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._load_pdf(path)
        elif suffix in [".docx", ".doc"]:
            return self._load_word(path)
        elif suffix in [".xlsx", ".xls"]:
            return self._load_excel(path)
        elif suffix == ".csv":
            return self._load_csv(path)
        elif suffix in [".md", ".markdown", ".txt"]:
            return self._load_text(path)
        else:
            raise DocumentLoaderError(f"不支持的文件格式: {suffix}")

    def _load_pdf(self, path: Path) -> str:
        """加载 PDF 文件，优先使用 PyMuPDF，降级到 pypdf"""
        # 尝试 PyMuPDF（效果更好）
        try:
            import fitz

            text_parts = []
            with fitz.open(str(path)) as doc:
                for page in doc:
                    text = page.get_text()
                    text_parts.append(text)
            return self._clean_text("\n".join(text_parts))
        except ImportError:
            pass
        except Exception as e:
            log.warning(f"PyMuPDF 解析失败，尝试 pypdf: {e}")

        # 降级到 pypdf
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return self._clean_text("\n".join(text_parts))
        except ImportError:
            raise DocumentLoaderError(
                "加载 PDF 需要安装 PyMuPDF 或 pypdf: " "pip install PyMuPDF 或 pip install pypdf"
            )
        except Exception as e:
            raise DocumentLoaderError(f"PDF 解析失败: {e}")

    def _load_word(self, path: Path) -> str:
        """加载 Word 文件"""
        try:
            import docx
        except ImportError:
            raise DocumentLoaderError("加载 Word 需要安装 python-docx: pip install python-docx")

        try:
            document = docx.Document(str(path))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

            # 提取表格内容
            tables_text = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables_text.append(" | ".join(cells))

            all_text = "\n\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n" + "\n".join(tables_text)

        except Exception as e:
            raise DocumentLoaderError(f"Word 解析失败: {e}")

        return self._clean_text(all_text)

    def _load_excel(self, path: Path) -> str:
        """加载 Excel 文件"""
        try:
            df = pd.read_excel(str(path))
            return self._dataframe_to_text(df)
        except Exception as e:
            raise DocumentLoaderError(f"Excel 解析失败: {e}")

    def _load_csv(self, path: Path) -> str:
        """加载 CSV 文件"""
        try:
            df = pd.read_csv(str(path))
            return self._dataframe_to_text(df)
        except Exception as e:
            raise DocumentLoaderError(f"CSV 解析失败: {e}")

    def _load_text(self, path: Path) -> str:
        """加载 Markdown / TXT 文件"""
        try:
            # 尝试 UTF-8
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # 降级到 GBK
            try:
                text = path.read_text(encoding="gbk")
            except UnicodeDecodeError as e:
                raise DocumentLoaderError(f"文件编码无法识别: {e}")

        return self._clean_text(text)

    def _dataframe_to_text(self, df: pd.DataFrame) -> str:
        """将 DataFrame 转换为文本"""
        # 去掉全空行和全空列
        df = df.dropna(how="all").dropna(axis=1, how="all")
        return self._clean_text(df.to_string(index=False))

    def _clean_text(self, text: str) -> str:
        """
        清洗文本
        - 合并连续空行
        - 去除行首行尾空白
        - 去除分页符等特殊字符
        """
        if not text:
            return ""

        # 替换常见特殊空白字符
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\f", "\n")  # 分页符

        # 去除每行首尾空白，但保留空行用于段落分隔
        lines = text.split("\n")
        cleaned_lines = [line.strip() for line in lines]

        # 合并连续空行
        result_lines = []
        prev_empty = False
        for line in cleaned_lines:
            is_empty = line == ""
            if is_empty and prev_empty:
                continue
            result_lines.append(line)
            prev_empty = is_empty

        # 去除首尾空行
        while result_lines and result_lines[0] == "":
            result_lines.pop(0)
        while result_lines and result_lines[-1] == "":
            result_lines.pop()

        return "\n".join(result_lines)


# 全局文档加载器实例
document_loader = DocumentLoader()
